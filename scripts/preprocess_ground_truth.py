"""Create ground-truth text files from the Word samples.

Parses the .docx files in samples/original/ (templates 1-8) and
samples/tables/ (templates 9-16) and writes one .txt per sample to
ground_truth/. samples/handwritting/ is intentionally skipped: it has the
same textual content as samples/original/, only the fonts differ.

Final output: 16 templates x 8 samples = 128 text files.
"""

import os
import zipfile
from lxml import etree
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _iter_block_text(paragraphs) -> list[str]:
    """Extract non-empty text from a list of paragraphs."""
    return [p.text for p in paragraphs if p.text.strip()]


def _cell_text(tc) -> str:
    """Join paragraphs within a cell with newlines, preserving internal line breaks."""
    para_texts = []
    for child in tc:
        if child.tag.split("}")[-1] == "p":
            run_text = "".join(t.text or "" for t in child.iter(f"{{{W_NS}}}t"))
            para_texts.append(run_text)
    return "\n".join(p for p in para_texts if p.strip())


def _table_texts(table) -> list[str]:
    """Extract cell texts directly from XML, avoiding python-docx phantom cell expansion."""
    texts = []
    for tr in table._tbl:
        if tr.tag.split("}")[-1] != "tr":
            continue
        for tc in tr:
            if tc.tag.split("}")[-1] != "tc":
                continue
            text = _cell_text(tc)
            if text:
                texts.append(text)
    return texts


def get_Word_text(path: str) -> str:
    doc = Document(path)
    parts = []

    # Headers (one per unique section header)
    seen_headers: set[str] = set()
    for section in doc.sections:
        header = section.header
        if header.is_linked_to_previous:
            continue
        header_text = "\n".join(_iter_block_text(header.paragraphs))
        if header_text and header_text not in seen_headers:
            parts.append(header_text)
            seen_headers.add(header_text)

    # Body: paragraphs and tables in document order
    for block in doc.element.body:
        tag = block.tag.split("}")[-1]
        if tag == "p":
            text = Paragraph(block, doc).text.strip()
            if text:
                parts.append(text)
        elif tag == "tbl":
            parts.extend(_table_texts(Table(block, doc)))

    # Footers (one per unique section footer)
    seen_footers: set[str] = set()
    for section in doc.sections:
        footer = section.footer
        if footer.is_linked_to_previous:
            continue
        footer_text = "\n".join(_iter_block_text(footer.paragraphs))
        if footer_text and footer_text not in seen_footers:
            parts.append(footer_text)
            seen_footers.add(footer_text)

    # Footnotes from word/footnotes.xml inside the docx zip
    with zipfile.ZipFile(path) as z:
        if "word/footnotes.xml" in z.namelist():
            root = etree.fromstring(z.read("word/footnotes.xml"))
            for fn in root.findall(f"{{{W_NS}}}footnote"):
                fn_id = fn.get(f"{{{W_NS}}}id")
                if fn_id in ("-1", "0"):   # skip separator/continuation notes
                    continue
                fn_text = "".join(
                    t.text or "" for t in fn.findall(f".//{{{W_NS}}}t")
                ).strip()
                if fn_text:
                    parts.append(f"[Footnote {fn_id}] {fn_text}")

    return "\n".join(parts)


BASE = os.path.join(os.path.dirname(__file__), "..")
# Source folders that contribute unique textual content. handwritting/ is
# excluded: same text as original/, different fonts only.
SOURCE_DIRS = [
    os.path.join(BASE, "samples", "original"),
    os.path.join(BASE, "samples", "tables"),
]
OUTPUT_DIR = os.path.join(BASE, "ground_truth")

os.makedirs(OUTPUT_DIR, exist_ok=True)


def main() -> None:
    count = 0
    for src_dir in SOURCE_DIRS:
        for fname in sorted(os.listdir(src_dir)):
            if not fname.endswith(".docx"):
                continue
            docx_path = os.path.join(src_dir, fname)
            txt_path = os.path.join(OUTPUT_DIR, fname.replace(".docx", ".txt"))
            if os.path.exists(txt_path):
                raise FileExistsError(
                    f"Name collision: {txt_path} already written from another source dir"
                )
            text = get_Word_text(docx_path)
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(text)
            count += 1
            print(f"Saved: {txt_path}")
    print(f"\nDone. Wrote {count} text files to {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
